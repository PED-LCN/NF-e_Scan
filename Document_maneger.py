import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor,Inches

def add_blank_lines(document, count):
    for _ in range(count):
        document.add_paragraph()

def add_formatted_paragraph(document, text):
    p = document.add_paragraph()
    run = p.add_run(text)
    font = run.font
    font.name = 'Courier New'
    font.size = Pt(16)
    run.bold = True
    font.color.rgb = RGBColor(0, 0, 0)
    return p

def create_document(dados_faturas, docx_file):
    if dados_faturas:
        document = Document()
        
        document.sections[0].top_margin = Inches(0.30)
        p_header = add_formatted_paragraph(document, 'EMPRESA: ETICAL ETIQUETAS CARUARU LTDA')
        p_header.style = document.styles['Heading 1']
        numero_global = dados_faturas[0]['numero'].split('/')[0] if dados_faturas else 'N/A'
        p_nota = add_formatted_paragraph(document, f'NOTA: {numero_global}')
        p_nota.alignment = WD_ALIGN_PARAGRAPH.LEFT
        add_formatted_paragraph(document, '-' * 44)
        add_blank_lines(document, 2)
        
        for i, item in enumerate(dados_faturas):
            if i > 0 and i % 4 == 0:
                document.add_page_break()
                p_header = add_formatted_paragraph(document, 'EMPRESA: ETICAL ETIQUETAS CARUARU LTDA')
                p_header.style = document.styles['Heading 1']
                p_nota = add_formatted_paragraph(document, f'NOTA: {numero_global}')
                p_nota.alignment = WD_ALIGN_PARAGRAPH.LEFT
                add_formatted_paragraph(document, '-' * 44)
                add_blank_lines(document, 2)
                
            add_formatted_paragraph(document, f"NOTA FISCAL: {item['numero']}")
            add_formatted_paragraph(document, f"DATA: {item['vencimento']}")
            add_formatted_paragraph(document, f"VALOR: {item['valor']}")
            
            if (i + 1) % 4 != 0 and i < len(dados_faturas) - 1:
                add_blank_lines(document, 2)

        document.save('documento_notas_fiscais.docx')
        print("Arquivo 'documento_notas_fiscais.docx' gerado com sucesso!")
    else:
        print("Nenhum dado de fatura encontrado. O arquivo Word não foi gerado.")